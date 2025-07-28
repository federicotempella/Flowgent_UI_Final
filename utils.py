import streamlit as st
import pandas as pd 
import base64
import json
import re
import openai
import datetime
import gspread
import fitz
import requests
import glob
from docx import Document
from bs4 import BeautifulSoup
from oauth2client.service_account import ServiceAccountCredentials
from PyPDF2 import PdfReader
from date import datetime

# utils.py

import pandas as pd
import re
# altri import se presenti...

import openai

def log_gpt_fallback(tipo, ruolo, industry, trigger, pain, kpi, note=""):
    try:
        scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
        creds = ServiceAccountCredentials.from_json_keyfile_dict(service_account_info, scope)
        client = gspread.authorize(creds)

        sheet = client.open("AI_SalesBot_UI_Log").worksheet("GPT_Fallback_Log")
        now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        user_id = st.session_state.get("user", "anonimo")

        sheet.append_row([
            now,
            user_id,
            tipo,
            ruolo,
            industry,
            trigger,
            ", ".join(pain if isinstance(pain, list) else [pain]),
            ", ".join(kpi if isinstance(kpi, list) else [kpi]),
            note
        ])
    except Exception as e:
        st.warning(f"❌ Errore nel log fallback GPT: {e}")
        
import requests
from bs4 import BeautifulSoup

def perform_deep_research(company: str, role: str = "", trigger: str = "") -> str:
    queries = [
        f"{company} sito ufficiale",
        f"{company} tecnologia o software in uso",
        f"{company} progetti digitali recenti o roadmap innovazione",
        f"{company} offerte lavoro {role}" if role else f"{company} offerte lavoro",
    ]

    # Trigger-based enrichment
    if trigger:
        if "ERP" in trigger or "migrazione" in trigger.lower():
            queries.append(f"{company} ERP in uso o migrazione SAP")
        elif "customer portal" in trigger.lower() or "self-service":
            queries.append(f"{company} portale clienti o canale digitale B2B")
        elif "AI" in trigger.lower() or "machine learning" in trigger.lower():
            queries.append(f"{company} progetti intelligenza artificiale")

    results = []

    # 🔍 1. Esegui una vera ricerca web (solo sul trigger, se presente)
    if trigger:
        try:
            search_query = f"{company} {trigger}".replace(" ", "+")
            headers = {"User-Agent": "Mozilla/5.0"}
            res = requests.get(f"https://www.bing.com/search?q={search_query}", headers=headers, timeout=5)
            if res.status_code == 200:
                soup = BeautifulSoup(res.text, "html.parser")
                first_result = soup.find("li", class_="b_algo")
                if first_result:
                    title = first_result.find("h2").get_text(strip=True)
                    snippet = first_result.find("p").get_text(strip=True)
                    link = first_result.find("a")["href"]
                    results.append(f"🌐 Approfondimento su “{trigger}” trovato online:\n🔗 {title}\n🧠 {snippet}\n🔗 {link}")
        except Exception as e:
            results.append(f"🌐 Errore nella ricerca web reale: {e}")

    # 🤖 2. Esegui ricerche simulate via GPT-4o per le altre query
    for query in queries:
        try:
            search_response = openai.ChatCompletion.create(
                model="gpt-4o",
                messages=[{
                    "role": "user",
                    "content": f"""Cerca online: {query}.
Dammi un estratto utile e sintetico (max 2 frasi) che aiuti a capire se l’azienda ha iniziative in corso o evoluzioni rilevanti."""
                }],
                temperature=0.4,
            )
            answer = search_response.choices[0].message.content.strip()
            results.append(f"🔍 {query} → {answer}")
        except Exception as e:
            results.append(f"❌ Errore su {query}: {e}")

    return "\n\n".join(results)

def generate_pain_from_trigger(trigger, ruolo):
    prompt = f"""Mi dai un esempio di pain point che potrebbe avere un {ruolo} se nota questo trigger: {trigger}? Rispondi in 1 frase."""
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"GPT error: {e}"

def generate_kpi_from_trigger(trigger, ruolo):
    prompt = f"""Quale KPI potrebbe essere impattato per un {ruolo} in presenza di questo trigger: {trigger}? Rispondi in 1 frase."""
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"GPT error: {e}"

def generate_pain_kpi_from_context(ruolo, industry, context=""):
    # Usa GPT per generare pain e KPI a partire dal ruolo e (se disponibile) il contesto Deep Research
    prompt = f"""
Sei un esperto B2B. In base a questo contesto aziendale:

--- CONTENUTO ---
{context}
------------------

Scrivi 2 pain point e 2 KPI che potrebbero essere rilevanti per un {ruolo} nel settore {industry}. Rispondi in elenco puntato, prima i pain poi i KPI.
"""
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.6,
        )
        text = response.choices[0].message.content.strip()
        pain = []
        kpi = []

        for line in text.splitlines():
            clean = line.strip("-• ").strip()
            if clean:
                if len(pain) < 2:
                    pain.append(clean)
                else:
                    kpi.append(clean)
        return pain, kpi
    except Exception as e:
        return [], []

def log_buyer_persona_submission(user_id, role, industry, pain, kpi, suggestion):
    scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
    creds = ServiceAccountCredentials.from_json_keyfile_name("client_secret.json", scope)
    client = gspread.authorize(creds)

    sheet = client.open("AI_SalesBot_UI_Log").worksheet("BuyerPersona_Log")

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    sheet.append_row([
        now,
        user_id,
        role,
        industry,
        ", ".join(pain),
        ", ".join(kpi),
        suggestion,
        "🟡 In attesa approvazione"
    ])


# 🔽 Inserisci subito dopo gli import
import json
import os

def load_persona_matrix_by_industry(industry):
    """
    Carica tutte le buyer persona per un dato settore da tutte le fonti (master, custom, live)
    """
    all_data = load_all_buyer_personas()
    matrix = {}

    for role, data in all_data.items():
        industries = data.get("industries", {})
        if industry in industries:
            matrix[role] = industries[industry]

    return matrix

# === Setup: credenziali da secrets ===
openai.api_key = st.secrets["OPENAI_API_KEY"]
assistant_id = st.secrets["ASSISTANT_ID"]
sheet_id = st.secrets["GOOGLE_SHEET_ID"]
service_account_info = json.loads(st.secrets["SERVICE_ACCOUNT_JSON"])

# === Setup: Google Sheets ===
def load_google_sheet():
    scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
    creds = ServiceAccountCredentials.from_json_keyfile_dict(service_account_info, scope)
    client = gspread.authorize(creds)
    return client.open_by_key(sheet_id)

def save_to_library(tipo, contenuto):
    try:
        sheet = load_google_sheet()
        library_tab = sheet.worksheet("library")
        library_tab.append_row([
            datetime.datetime.now().isoformat(),
            tipo,
            contenuto,
            st.session_state.get("user", "Anonimo")
        ])
        st.success("✅ Contenuto salvato nella tua libreria.")
    except Exception as e:
        st.error(f"Errore nel salvataggio in libreria: {str(e)}")


# === Branding ===
def load_logo(path="logo.png"):
    try:
        with open(path, "rb") as f:
            return f.read()
    except FileNotFoundError:
        st.warning("⚠️ Logo non trovato. Verrà usato un placeholder.")
        return None

# === 📁 Upload & Parsing Excel ===
def load_uploaded_excel(uploaded_file):
    df = pd.read_excel(uploaded_file)
    return df

def parse_uploaded_file(uploaded_file):
    df = load_uploaded_excel(uploaded_file)
    df = df.dropna(how='all')  # rimuove righe vuote

    # Validazione colonne
    expected_cols = [
        "Name", "Company", "Role", "LinkedIn", 
        "Manually Found Trigger - Something you read / inferred", 
        "Manually Found Trigger - Interesting LinkedIn Post", 
        "Manually Found Trigger - LinkedIn Signal", 
        "Manually Found Trigger - Contact in Common Relevant", 
        "Manually Found Trigger - Company Signal"
    ]
    for col in expected_cols:
        if col not in df.columns:
            raise ValueError(f"Colonna mancante: {col}")

    # Gestione trigger: concatenazione in lista
    def extract_triggers(row):
        triggers = []
        for col in expected_cols[4:]:
            value = row.get(col)
            if pd.notna(value):
                split_values = [v.strip() for v in str(value).split(";") if v.strip()]
                triggers.extend(split_values)
        return triggers

    df["Triggers"] = df.apply(extract_triggers, axis=1)
    return df

def parse_excel_file(uploaded_file):
    return parse_uploaded_file(uploaded_file)

from PyPDF2 import PdfReader

def parse_pdf_files(pdf_files):
    parsed_results = {}
    for pdf in pdf_files:
        try:
            reader = PdfReader(pdf)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            parsed_results[pdf.name] = text.strip()
        except Exception as e:
            parsed_results[pdf.name] = f"Errore nel parsing: {e}"
    return parsed_results

# === 🎯 CAMPAGNA ===
def start_campaign_flow(parsed_df=None):
    if parsed_df is not None:
        st.success(f"{len(parsed_df)} contatti caricati correttamente per la campagna.")
        st.dataframe(parsed_df[["Name", "Company", "Role", "Triggers"]])
    else:
        st.warning("📂 Nessun file caricato. Carica un file Excel per avviare la campagna.")

# === 🤖 SIMULATORE ===
def simulate_conversation():
    st.subheader("🤖 Simulatore GPT")
    user_input = st.text_area("Scrivi un prompt...")
    if st.button("Invia"):
        if user_input:
            response = openai.ChatCompletion.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": user_input}],
                temperature=0.7,
            )
            st.write("**Risposta GPT:**", response.choices[0].message.content)
        else:
            st.warning("Inserisci un prompt prima di inviare.")

def generate_multichannel_sequence(contact_df, sequence_type="AE", use_inmail=False):
    sequences = []

    for _, row in contact_df.iterrows():
        name = row.get("Name", "Contatto")
        company = row.get("Company", "Azienda")
        role = row.get("Role", "Ruolo")
        trigger = row.get("Trigger combinato", "")

        step_list = []

        if sequence_type == "AE":
            step_list = [
                "[Giorno 1] ✉️ Email di apertura (TIPPS + COI se score >= 4)",
                "[Giorno 3] 🤝 Richiesta connessione LinkedIn (senza nota)",
                "[Giorno 5] 💬 DM LinkedIn personalizzato (se accetta)",
                "[Giorno 6] 📎 Follow-up email con case study", 
                "[Giorno 8] 💬 DM video / voice note (se non risponde)",
                "[Giorno 10] ✉️ Email bump consultiva",
                "[Giorno 12] 💬 Ultimo DM: chiusura + soft CTA",
                "[Giorno 14] 🧠 Invia asset finale (es. guida, template)"
            ]
            if use_inmail:
                step_list.insert(2, "[Giorno 4] ✉️ InMail teaser (se ignorato)")
        else:  # SDR lunga 11 step
            step_list = [
                "[Giorno 1] ✉️ Email teaser (1-riga, aggancia problema)",
                "[Giorno 2] 🤝 Connessione LinkedIn (senza nota)",
                "[Giorno 3] ✉️ InMail: hook provocatorio + mini payoff",
                "[Giorno 4] 💬 DM breve: 'Curioso come lo affrontate internamente?'",
                "[Giorno 5] 🎙️ Voice Note LinkedIn – max 30 sec, friendly",
                "[Giorno 6] ✉️ Follow-up Email (TIPPS)",
                "[Giorno 7] 📞 Cold Call (se apertura o trigger forte)",
                "[Giorno 9] 💬 Interazione post LinkedIn (like/commento)",
                "[Giorno 11] ✉️ Bump: 'Mando tutto a cestino o ha senso riprendere?'",
                "[Giorno 13] 📎 DM consultivo: case simile + spunto",
                "[Giorno 15] 🏁 Ultimo touch o passa ad AE"
            ]

        # Inserimento automatico CALL o BUMP
        call_trigger_keywords = ["interesse", "ebook", "pdf", "evento", "demo", "profilo", "engagement", "aperto", "letto"]
        bump_keywords = ["nessuna risposta", "ghosting", "non ha risposto", "nessun feedback", "ignorato"]

        if any(k in trigger.lower() for k in call_trigger_keywords):
            step_list.append("[Giorno 16] 📞 Call suggerita – il contatto ha mostrato segnali di interesse")
        elif any(k in trigger.lower() for k in bump_keywords):
            step_list.append("[Giorno 17] ⏪ Bump finale (DM o Email) – stimola reazione se silenzio prolungato")

        sequences.append({
            "Name": name,
            "Company": company,
            "Role": role,
            "Sequenza multicanale": step_list
        })

    return sequences

# === 📬 POST GENERATOR ===
def generate_post():
    st.subheader("✍️ Generatore di contenuti")
    idea = st.text_input("Idea o tema del post:")
    tone = st.selectbox("Tono", ["Professionale", "Informale", "Provocatorio"])
    if st.button("Genera post"):
        if idea:
            response = openai.ChatCompletion.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": f"Genera un post LinkedIn in tono {tone}."},
                    {"role": "user", "content": idea}
                ],
                temperature=0.8,
            )
            post = response.choices[0].message.content
            st.text_area("Post generato", post, height=200)
            if st.button("💾 Salva in libreria"):
                save_to_library("Post LinkedIn", post)
                st.success("Post salvato nella tua libreria.")
        else:
            st.warning("Inserisci un'idea per generare il post.")

# === 🔁 AGGIORNAMENTI MODELLO ===
def check_for_updates():
    try:
        sheet = load_google_sheet()
        worksheet = sheet.worksheet("Model_Updates")
        values = worksheet.get_all_values()
        if len(values) <= 1:
            st.info("Nessun aggiornamento disponibile.")
            return
        headers = values[0]
        records = values[1:]
        st.subheader("🔄 Aggiornamenti al modello")
        for i, row in enumerate(records, start=2):  # partendo dalla riga 2
            update = dict(zip(headers, row))
            st.markdown(f"**🆕 {update['Titolo']}** – {update['Descrizione']}")
            col1, col2 = st.columns(2)
            with col1:
                if st.button(f"✅ Accetta {update['Titolo']}", key=f"accept_{i}"):
                    worksheet.update_cell(i, headers.index("Accettato_colonna") + 1, "Sì")
                    st.success(f"{update['Titolo']} accettato.")
            with col2:
                if st.button(f"❌ Ignora {update['Titolo']}", key=f"ignore_{i}"):
                    worksheet.update_cell(i, headers.index("Accettato_colonna") + 1, "No")
    except Exception as e:
        st.error(f"Errore nella lettura degli aggiornamenti: {str(e)}")

# === 📅 DAILY TASKS ===
def show_daily_tasks():
    try:
        sheet = load_google_sheet()
        worksheet = sheet.worksheet("UI_Log")
        values = worksheet.get_all_values()
        if len(values) <= 1:
            st.info("Nessuna attività registrata oggi.")
            return
        headers = values[0]
        records = values[1:]
        df = pd.DataFrame(records, columns=headers)
        df['Timestamp'] = pd.to_datetime(df['Timestamp'], errors='coerce')
        today = datetime.datetime.now().date()
        filtered = df[df['Timestamp'].dt.date == today]
        st.subheader("📅 Le tue azioni recenti")
        st.dataframe(filtered)
    except Exception as e:
        st.error(f"Errore nella lettura delle attività: {str(e)}")

# === 📋 FEEDBACK FORM ===
def show_feedback_form():
    st.subheader("📝 Lascia un feedback")
    feedback = st.text_area("Il tuo feedback...")
    if st.button("Invia feedback"):
        if feedback:
            sheet = load_google_sheet()
            feedback_tab = sheet.worksheet("UI_Log")
            feedback_tab.append_row([
                datetime.datetime.now().isoformat(),
                st.session_state.get("user", "Anonimo"),
                st.session_state.get("email", ""),
                st.session_state.get("role", ""),
                st.session_state.get("level", ""),
                st.session_state.get("lang", ""),
                "Feedback",
                feedback,
                "", "", ""
            ])
            st.success("Feedback inviato!")
        else:
            st.warning("Scrivi qualcosa prima di inviare.")

# === 💡 SETTINGS ===
def show_settings(): 
    st.subheader("⚙️ Modifica le tue impostazioni")

    user = st.text_input("Nome utente (facoltativo)", value=st.session_state.get("user", ""))
    level = st.selectbox("Livello", ["Beginner", "Intermediate", "Advanced"], index=["Beginner", "Intermediate", "Advanced"].index(st.session_state.get("level", "Beginner")))
    lang = st.selectbox("Lingua", ["Italiano", "English"], index=["Italiano", "English"].index(st.session_state.get("lang", "Italiano")))

    if st.button("Salva impostazioni"):
        st.session_state["user"] = user if user else "Anonimo"
        st.session_state["level"] = level
        st.session_state["lang"] = lang
        st.success("Impostazioni aggiornate.")

# === 📚 LIBRERIA ===
def show_library():
    try:
        sheet = load_google_sheet()
        worksheet = sheet.worksheet("library")
        values = worksheet.get_all_values()
        if len(values) <= 1:
            st.info("La libreria è vuota.")
            return
        headers = values[0]
        records = values[1:]
        df = pd.DataFrame(records, columns=headers)
        st.subheader("📚 La tua libreria")
        st.dataframe(df[["Timestamp", "Tipo", "Contenuto"]])
    except Exception as e:
        st.error(f"Errore nella lettura della libreria: {str(e)}")

# === 📊 REPORTS ===
def show_reports():
    sheet = load_google_sheet()
    st.subheader("📊 Log attività")
    for tab_name in ["UI_Log", "Main_Log"]:
        st.markdown(f"### {tab_name}")
        try:
            worksheet = sheet.worksheet(tab_name)
            values = worksheet.get_all_values()
            if len(values) <= 1:
                st.info(f"Nessun dato disponibile in {tab_name}.")
                continue
            headers = values[0]
            records = values[1:]
            df = pd.DataFrame(records, columns=headers)
            st.dataframe(df)
        except Exception as e:
            st.error(f"Errore nella lettura del foglio {tab_name}: {str(e)}")

# === 🔒 PRIVACY POLICY ===
def show_privacy_policy():
    st.subheader("📄 Data Privacy & Proprietà Intellettuale")
    with open("guide_placeholder.pdf", "rb") as f:
        st.download_button("📥 Scarica policy completa", f, file_name="guide_placeholder.pdf")

# === 🏠 SCREEN 0 ===
def show_screen_zero():
    st.markdown("### 👋 Benvenuto!")
    st.markdown("Questa è la schermata iniziale del Sales Bot. Usa la toolbar a destra per iniziare.")

# === 📊 Trigger → KPI → Messaggio suggerito ===
import pandas as pd
import re
import openai

def analyze_triggers_and_rank(df, parsed_pdf=None, manual_input=None, buyer_personas=None):
    if buyer_personas is None:
        buyer_personas = {}

    def extract_known_triggers(text):
        triggers = []
        for role_data in buyer_personas.values():
            for industry_data in role_data.get("industries", {}).values():
                for p in industry_data.get("pain", []):
                    if p and re.search(re.escape(p), text, re.IGNORECASE):
                        triggers.append(p)
        return list(set(triggers))

    # Unisci testo da PDF e input manuale in un solo corpus per analisi
    combined_texts = []
    if parsed_pdf:
        for content in parsed_pdf.values():
            combined_texts.append(content)
    if manual_input:
        combined_texts.append(manual_input)

    full_context = " ".join(combined_texts).lower()
    results = []

    for _, row in df.iterrows():
        name = row.get("Name", "N/A")
        company = row.get("Company", "N/A")
        role = row.get("Role", "N/A")

        # 🧠 Ricerca approfondimenti (Deep Research) se abilitata
        deep_note = ""
        if st.session_state.get("deep_research", False) and company != "N/A":
            try:
                deep_note = perform_deep_research(company=company, role=role, trigger=row.get("Triggers", ""))
            except Exception as e:
                deep_note = f"Errore deep research: {e}"

        trigger_cell = row.get("Triggers", "")
        found_triggers = []

        # Analisi Excel triggers (separati da virgola o punto e virgola)
        for t in re.split(r",|;", str(trigger_cell)):
            t_clean = t.strip().lower()
            if t_clean:
                found_triggers.append(t_clean)

        # Arricchimento: aggiungi anche i trigger trovati in PDF/AI notes
        extra_triggers = extract_known_triggers(full_context) if full_context else []
        all_triggers = [t.lower() for t in found_triggers + extra_triggers if t]

        # Se non ci sono trigger trovati, utilizza un pain point predefinito (fallback) dal buyer_personas
        if not all_triggers:
            industry_data = None
            if role in buyer_personas:
                # prova a prendere il settore selezionato oppure il primo disponibile
                industries = buyer_personas[role].get("industries", {})
                sel_industry = st.session_state.get("selected_industry")
                if sel_industry and sel_industry in industries:
                    industry_data = industries.get(sel_industry, {})
                else:
                    # prendi il primo settore disponibile
                    for ind_data in industries.values():
                        industry_data = ind_data
                        break
            if industry_data:
                pain_list = industry_data.get("pain", [])
                if pain_list:
                    first_pain = pain_list[0]
                    all_triggers = [first_pain.lower()]
                    found_triggers = [first_pain.lower()]
                    trigger_cell = first_pain
            else:
                trigger_cell = ""

        # Se abbiamo trovato almeno un trigger (o impostato un fallback), calcola lo score e abbina pain/KPI noti
        score = 0
        pain_list = []
        kpi_list = []
        framework = ""
        suggestion = ""
        if all_triggers:
            # Calcola punteggio di rilevanza basato sui trigger trovati
            score = len(all_triggers)

            # Identifica pain e KPI corrispondenti al trigger principale
            primary_trigger = all_triggers[0]
            for role_data in buyer_personas.values():
                for industry_data in role_data.get("industries", {}).values():
                    for pain_point, kpi_val, sugg, fw in zip(
                        industry_data.get("pain", []),
                        industry_data.get("kpi", []),
                        industry_data.get("suggestion", []),
                        industry_data.get("framework", []),
                    ):
                        if pain_point and primary_trigger and re.search(re.escape(pain_point), primary_trigger, re.IGNORECASE):
                            pain_list.append(pain_point)
                            if kpi_val:
                                kpi_list.append(kpi_val)
                            if sugg:
                                suggestion = sugg
                            if fw:
                                framework = fw

        # Pulisce e rimuove duplicati
        pain_str = ", ".join(list(set(pain_list)))
        kpi_str = ", ".join(list(set(kpi_list)))

        # Costruzione finale del record
        results.append({
            "Name": name,
            "Company": company,
            "Role": role,
            "Trigger combinato": ", ".join(all_triggers),
            "Score": score,
            "KPI consigliati": kpi_str,
            "Framework suggerito": framework,
            "Suggerimento di messaggio": suggestion,
            "Note Deep": deep_note
        })

    return pd.DataFrame(results)

# === 🧠 Personalizzazione multivariabile con GPT ===
import openai
import pandas as pd
import json
import os

# Funzione per caricare framework live
def load_frameworks():
    path_master = "resources/frameworks_master.json"
    if os.path.exists(path_master):
        with open(path_master, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def generate_personalized_messages(ranked_df, framework_override=None, frameworks_all=None):
    import openai
    import pandas as pd

    messages = []

    for _, row in ranked_df.iterrows():
        row = fallback_missing_fields(row.copy())

        name = row.get("Name", row.get("Nome", ""))
        company = row.get("Company", row.get("Azienda", ""))
        role = row.get("Role", row.get("Ruolo", ""))
        triggers = row.get("Trigger combinato", row.get("Trigger", ""))
        note_deep = row.get("Note Deep", "")
        framework = framework_override or row.get("Framework suggerito", "TIPPS")
        kpi = row.get("KPI impattati", "")
        pain = row.get("Pain Point", "")
        industry = row.get("Settore", "custom")
        
        context = f"""Il contatto {name} lavora in {company} come {role}.
    Trigger rilevati: {triggers}.
    Insight aggiuntivi da ricerche online:
    {note_deep}
    """
        # Aggiunta sintomo se disponibile
        symptom = row.get("Symptom", "")
        if symptom:
            context += f"\nSintomi osservati: {symptom}"

        prompt = f"""Agisci come un SDR esperto.
Usa il framework {framework} per scrivere un primo messaggio di contatto.
Contesto:
{context}
Scrivi in modo sintetico e d’impatto."""

    # Chiamata GPT vera
        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.6,
        )
        message = response.choices[0].message.content.strip()
        messages.append({
            "Name": name,
            "Company": company,
            "Role": role,
            "Trigger combinato": triggers,
            "Message": message,
            "Framework": framework
        })
        
    return pd.DataFrame(messages)

        # 🔄 Se alcuni campi mancano, prova a completarli con GPT o dati persona
        if not pain and trigger and ruolo:
            pain = generate_pain_from_trigger(trigger, ruolo)
        if not kpi and trigger and ruolo:
            kpi = generate_kpi_from_trigger(trigger, ruolo)
        symptom = row.get("Symptom", "")
        if not symptom and pain and ruolo:
            symptom = generate_symptom_from_pain(pain, ruolo)

        # 📋 Logging fallback GPT se abbiamo riempito campi mancanti
        if not row.get("Pain Point") or not row.get("KPI impattati"):
            log_gpt_fallback(
                tipo="Completamento GPT",
                ruolo=ruolo,
                industry=industry,
                trigger=trigger,
                pain=pain,
                kpi=kpi,
                note="Completato perché mancava pain/kpi nella riga"
            )

        # 🔁 Gestione framework (override o auto da score)
        framework_id = row.get("Framework suggerito", "")
        if framework_override and framework_override != "Auto (da score)":
            framework_id = framework_override

        # Carica frameworks se non già fornito
        if frameworks_all is None:
            frameworks_all = load_all_frameworks()
        selected_fw = frameworks_all.get(framework_id) if frameworks_all else None
        if not selected_fw:
            # Tentativo fallback in base al settore (se definito)
            fallback_fw = get_default_framework_by_industry(industry)
            if fallback_fw:
                selected_fw = fallback_fw
                framework_id = fallback_fw.get("id", "TIPPS")

        # Estrai nome, descrizione e struttura del framework selezionato
        framework_name = selected_fw.get("name", "") if selected_fw else ""
        description = selected_fw.get("description", "") if selected_fw else ""
        structure_text = "\n".join([f"- {step}" for step in selected_fw.get("structure", selected_fw.get("rules", []))]) if selected_fw else ""

        # ✍️ Costruzione del prompt per GPT
        symptom_line = f"- Sintomo: {symptom}\n" if symptom else ""
        prompt = f"""Scrivi un messaggio di sales outbound (email o LinkedIn) basato su questo contesto:

📘 Framework: {framework_name}
📋 Descrizione: {description}

📎 Contesto:
- Nome: {nome}
- Azienda: {azienda}
- Ruolo: {ruolo}
- Trigger: {trigger}
- KPI: {kpi}
- Pain: {pain}
{symptom_line}📚 Approfondimenti aggiuntivi:
{extra_notes}

🎯 Obiettivo: ottenere risposta o apertura.
🗣️ Tono: diretto, rilevante e professionale.
✂️ Linguaggio: sintetico, rilevante, payoff chiaro.
❌ Evita frasi generiche.
✅ Concludi con una call-to-action soft.
"""

        # 🤖 Chiamata al modello GPT per generare il messaggio
        try:
            completion = openai.ChatCompletion.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.6,
            )
            message_text = completion.choices[0].message.content.strip()
        except Exception as e:
            # Usa un suggerimento predefinito come fallback se disponibile
            message_text = row.get("Suggerimento di messaggio") or f"(GPT fallito) {suggestion}" if 'suggestion' in locals() else f"Errore generazione: {e}"

        # Salva il messaggio generato e il contesto associato
        messages.append({
            "Nome": nome,
            "Azienda": azienda,
            "Ruolo": ruolo,
            "Pain Point": pain,
            "KPI impattati": kpi,
            "Messaggio generato": message_text,
            "Note Deep": extra_notes
        })

    return pd.DataFrame(messages)


# utils.py

import streamlit as st
import json, os

base_file = "resources/buyer_personas_master.json"
custom_file = "resources/buyer_personas.json"

def load_all_buyer_personas():
    base_data = {}
    custom_data = {}
    live_data = st.session_state.get("buyer_personas_live", {})

    if os.path.exists(base_file):
        with open(base_file, "r", encoding="utf-8") as f:
            base_data = json.load(f)

    if os.path.exists(custom_file):
        with open(custom_file, "r", encoding="utf-8") as f:
            custom_data = json.load(f)

    # Merge con priorità: live > custom > base
    merged = base_data.copy()

    for role, role_data in custom_data.items():
        if role not in merged:
            merged[role] = role_data
        else:
            merged[role]["industries"] = {
                **merged[role].get("industries", {}),
                **role_data.get("industries", {})
            }

    for role, role_data in live_data.items():
        if role not in merged:
            merged[role] = role_data
        else:
            merged[role]["industries"] = {
                **merged[role].get("industries", {}),
                **role_data.get("industries", {})
            }

    return merged


def fallback_missing_fields(row):
    bp_data = load_all_buyer_personas()
    ruolo = row.get("Ruolo", "")
    industry = "custom"  # puoi migliorarlo leggendo la colonna se presente

    persona = bp_data.get(ruolo, {}).get("industries", {}).get(industry, {})
    ruolo_simile = None

    if not row.get("Pain Point") and persona.get("pain"):
        row["Pain Point"] = persona["pain"][0]

    if not row.get("KPI impattati") and persona.get("kpi"):
        row["KPI impattati"] = persona["kpi"][0]

    

     if not row.get("Pain Point"):
        # Prova a cercare ruolo simile
        for candidate_role, data in bp_data.items():
            if candidate_role.lower() != ruolo.lower():
                candidate = data.get("industries", {}).get(industry, {})
                if candidate.get("pain"):
                    row["Pain Point"] = candidate["pain"][0]
                    ruolo_simile = candidate_role
                    break

    if not row.get("KPI impattati"):
        for candidate_role, data in bp_data.items():
            if candidate_role.lower() != ruolo.lower():
                candidate = data.get("industries", {}).get(industry, {})
                if candidate.get("kpi"):
                    row["KPI impattati"] = candidate["kpi"][0]
                    ruolo_simile = candidate_role
                    break
    

    if ruolo_simile:
        log_gpt_fallback(
            tipo="Suggerimento ruolo simile",
            ruolo=ruolo,
            industry=industry,
            trigger=row.get("Trigger rilevato", ""),
            pain=row.get("Pain Point", ""),
            kpi=row.get("KPI impattati", ""),
            note=f"Template usato da ruolo simile: {ruolo_simile}"
        )
  
    return row

def save_all_buyer_personas(data):
    # Aggiunge automaticamente i nuovi campi se mancanti
    for ruolo, role_data in data.items():
        industries = role_data.get("industries", {})
        for industry, industry_data in industries.items():
            if "symptom" not in industry_data:
                industry_data["symptom"] = []
            if "damage" not in industry_data:
                industry_data["damage"] = []

    # Salva nel file custom (quelle modificate/aggiunte dall’utente)
    with open("resources/buyer_personas.json", "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


import glob
from docx import Document
from PyPDF2 import PdfReader

def parse_persona_documents(directory_path: str, personas_dict: dict) -> dict:
    """
    Legge tutti i file .docx e .pdf nella cartella specificata e integra i dati estratti nel dizionario personas_dict.
    Ritorna il dizionario aggiornato.
    """
    # Per ogni file .docx o .pdf nella cartella
    for filepath in glob.glob(f"{directory_path}/*"):
        filename = filepath.split("/")[-1]
        if filename.lower().endswith(".docx"):
            doc = Document(filepath)
            # Determina l'industry dal titolo o dal nome file
            industry = None
            if doc.paragraphs:
                title_text = doc.paragraphs[0].text.strip()
                # Ci si aspetta che il titolo contenga l'industry dopo un trattino. Esempio: "Buyer Persona – Retail"
                if "–" in title_text:
                    industry = title_text.split("–")[-1].strip()
                else:
                    # Se non presente nel titolo, usa parte del nome file come fallback
                    industry = filename.replace(".docx", "").replace("buyer persona", "").replace("buyer personas", "").strip().capitalize()
            else:
                industry = filename.replace(".docx", "").capitalize()
            
            # Scansione del documento per ruoli e relativi punti
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            i = 0
            while i < len(paragraphs):
                para = paragraphs[i]
                # Individua l'inizio di una sezione per un ruolo (paragrafo in grassetto nel doc originale)
                # Possibile elenco di ruoli noti (CIO, IT Director, EDI Manager, ecc.), altrimenti assumiamo che 
                # un paragrafo senza due punti e non una delle intestazioni generali sia un ruolo
                if para.endswith(":") or para.lower() in ["top 3 pain points", "what they actually experience", "business damage", "3 tactical bis benefits"]:
                    # Skip intestazioni generali (che nel doc appaiono come testo normale ma in grassetto)
                    i += 1
                    continue
                # Heuristic: consideriamo un paragrafo come ruolo se è breve (<=5 parole) e inizia con lettera maiuscola
                # ed è noto oppure segue le intestazioni generali
                if 0 < len(para.split()) <= 5 and para.isupper() or para.istitle():
                    role = para
                    # Inizializza struttura se non presente
                    if role not in personas_dict:
                        personas_dict[role] = {"industries": {}}
                    if industry not in personas_dict[role]["industries"]:
                        personas_dict[role]["industries"][industry] = {
                            "pain": [], "symptom": [], "damage": [], 
                            "kpi": [], "suggestion": ""
                        }
                    # Raccoglie i 3 pain points successivi (indicati con "1. ", "2. ", "3. ")
                    pains = []
                    j = i + 1
                    while j < len(paragraphs) and len(pains) < 3:
                        if paragraphs[j][0].isdigit():  # inizia con numero (1., 2., ...)
                            pains.append(paragraphs[j].split(maxsplit=1)[1])  # prendi il testo dopo "1. "
                        else:
                            break
                        j += 1
                    # Raccoglie i 3 "What They Experience" (linee con "- ")
                    symptoms = []
                    while j < len(paragraphs) and len(symptoms) < 3:
                        if paragraphs[j].startswith("-"):
                            symptoms.append(paragraphs[j][2:])  # testo dopo "- "
                        else:
                            break
                        j += 1
                    # Raccoglie i 3 "Business Damage" (altre 3 linee con "- ")
                    damages = []
                    while j < len(paragraphs) and len(damages) < 3:
                        if paragraphs[j].startswith("-"):
                            damages.append(paragraphs[j][2:])
                        else:
                            break
                        j += 1
                    # Raccoglie i 3 "Tactical BIS Benefits" (indicati con "1. ", "2. ", "3." dopo i bullet)
                    benefits = []
                    while j < len(paragraphs) and len(benefits) < 3:
                        # Potrebbero riprendere la numerazione da 1 nuovamente
                        if paragraphs[j][0].isdigit():
                            benefits.append(paragraphs[j].split(maxsplit=1)[1])
                        else:
                            break
                        j += 1
                    # Integra i risultati nel dizionario persona
                    industry_entry = personas_dict[role]["industries"][industry]
                    for p in pains:
                        if p not in industry_entry["pain"]:
                            industry_entry["pain"].append(p)
                    for s in symptoms:
                        if s not in industry_entry["symptom"]:
                            industry_entry["symptom"].append(s)
                    for d in damages:
                        if d not in industry_entry["damage"]:
                            industry_entry["damage"].append(d)
                    # Per suggestion, se non impostato e ci sono benefits, usa il primo benefit come suggerimento
                    if industry_entry.get("suggestion") in [None, ""]:
                        if benefits:
                            industry_entry["suggestion"] = benefits[0]
                    # NOTA: i benefits aggiuntivi potrebbero essere conservati altrove, ad esempio in un campo separato se necessario.
                    # Aggiorna l'indice principale al termine di questa sezione ruolo
                    i = j
                    continue  # passa al prossimo ciclo senza incrementare i manualmente
                # Se il paragrafo corrente non rappresenta un ruolo, passa al successivo
                i += 1
        
        elif filename.lower().endswith(".pdf"):
            # Parsing PDF: estrai testo grezzo e poi usa eventualmente l'AI per identificare le parti chiave
            reader = PdfReader(filepath)
            full_text = ""
            for page in reader.pages:
                full_text += page.extract_text() + "\n"
            # Potenziale implementazione: usare l'AI per estrarre campi da full_text se formato non strutturato
            # Esempio (pseudo-codice):
            # prompt = f"Analizza il seguente testo e identifica per ogni ruolo i pain points, i sintomi, i danni e i benefici:\n{full_text}"
            # response = openai.ChatCompletion.create(..., messages=[...], ...)
            # parsed_info = response[...]  # dovrebbe restituire una struttura analizzabile (JSON string o simile)
            # personas_dict = merge_parsed_info(personas_dict, parsed_info)
            # Per semplicità, qui possiamo omettere la logica dettagliata e supporre che i PDF seguano lo stesso schema testuale dei docx.
            continue  # (Rimuovere o sostituire con la logica di parsing AI per PDF)
    return personas_dict

from docx import Document

def extract_symptom_and_damage_from_text(text: str) -> tuple[list[str], list[str]]:
    """Estrae symptom e damage dal testo libero."""
    symptoms, damages = [], []
    if not text:
        return symptoms, damages

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "Estrai dai seguenti appunti i sintomi (indizi del problema) e i danni (costi, rallentamenti, perdite)."},
                {"role": "user", "content": text}
            ],
            temperature=0.4,
        )
        answer = response.choices[0].message.content.strip()
        blocks = answer.split("\n")
        for line in blocks:
            if line.lower().startswith("symptom") or "sintomi" in line.lower():
                symptoms = [x.strip("•- ") for x in line.split(":")[1].split(",")]
            elif line.lower().startswith("damage") or "danno" in line.lower():
                damages = [x.strip("•- ") for x in line.split(":")[1].split(",")]
    except Exception as e:
        print(f"[⚠️] Errore estrazione symptom/damage: {e}")

    return symptoms, damages


