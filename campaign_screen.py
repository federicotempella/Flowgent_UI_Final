import streamlit as st
import pandas as pd
import io
from docx import Document
from io import BytesIO
import matplotlib.pyplot as plt 

from utils import (
    parse_excel_file,
    parse_pdf_files,
    analyze_triggers_and_rank,
    generate_personalized_messages,
    generate_multichannel_sequence,
    save_to_library,
    load_all_buyer_personas
)

st.subheader("🚀 Avvia una nuova campagna")

# 1. Caricamento Excel contatti

with st.expander("🔍 Deep Research (opzionale)"):
    deep_research_enabled = st.checkbox("Attiva Deep Research per ogni contatto")
    st.session_state["deep_research"] = deep_research_enabled

with st.expander("📁 1. Carica il file Excel contatti", expanded=True):
    uploaded_file = st.file_uploader("Seleziona il file Excel (.xlsx)", type=["xlsx"])
    if uploaded_file:
        try:
            df = parse_excel_file(uploaded_file)
            st.session_state["excel_df"] = df
            st.success(f"✅ {len(df)} contatti caricati correttamente.")
            st.dataframe(df[["Name", "Company", "Role", "Triggers"]])
        except Exception as e:
            st.error(f"❌ Errore durante il parsing del file: {e}")
    else:
        st.info("Carica un file Excel per visualizzare i contatti.")

# 2. Caricamento multiplo PDF
with st.expander("📄 2. Carica i PDF da usare come input (opzionale)", expanded=True):
    uploaded_pdfs = st.file_uploader("Carica uno o più file PDF", type=["pdf"], accept_multiple_files=True)
    if uploaded_pdfs:
        total_size = sum([f.size for f in uploaded_pdfs]) / (1024 * 1024)
        if len(uploaded_pdfs) > 10:
            st.warning("⚠️ Hai caricato più di 10 PDF. Consigliato: max 10.")
        if total_size > 50:
            st.warning(f"⚠️ Dimensione totale: {total_size:.1f}MB.")
        st.success(f"📄 {len(uploaded_pdfs)} PDF caricati.")
        if st.button("📌 Memorizza questi PDF"):
            if "pdf_memory" not in st.session_state:
                st.session_state["pdf_memory"] = []
            st.session_state["pdf_memory"].extend(uploaded_pdfs)
            st.success("PDF aggiunti alla memoria temporanea.")

# Parsing PDF
pdfs_mem = st.session_state.get("pdf_memory", [])
parsed_pdf = st.session_state.get("parsed_pdf", {})
if pdfs_mem and st.button("🧐 Elabora PDF ora"):
    parsed_pdf = parse_pdf_files(pdfs_mem)
    st.session_state["parsed_pdf"] = parsed_pdf
    for filename, content in parsed_pdf.items():
        with st.expander(f"📄 {filename}", expanded=False):
            st.markdown("**Testo estratto (parziale):**")
            st.write(content[:1500] + "..." if len(content) > 1500 else content)

# Industry & ruolo
industry = st.selectbox("Scegli il settore", ["automotive", "fashion retail", "CPG", "tier 1 automotive"])
buyer_personas = load_all_buyer_personas()
roles = list(buyer_personas.keys())
selected_role = st.selectbox("🌟 Seleziona un ruolo", roles)
st.markdown("### 💬 Chatta con l’AI per arricchire il contesto (facoltativo ma potente)")
user_prompt = st.text_area("Scrivi qui una domanda, carica contenuti o chiedi aiuto...", key="campaign_chat")

if st.button("✉️ Invia alla chat AI"):
    if user_prompt:
        with st.spinner("💬 Elaborazione in corso…"):
            response = openai.ChatCompletion.create(
                model="gpt-4o",
                messages=[
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.6,
            )
            ai_notes = response.choices[0].message.content
            st.markdown("**🧠 Risposta AI:**")
            st.write(ai_notes)
            st.session_state["ai_notes"] = ai_notes
    else:
        st.warning("Scrivi qualcosa prima di inviare.")

manual_input = st.session_state.get("ai_notes", "")
with st.expander("🧠 Potenzia l’input manuale"):
    col1, col2 = st.columns(2)
    if col1.button("🔎 Cerca sul web", key="web_manual_input"):
        st.session_state["web_manual_input"] = perform_web_search(manual_input)
        st.info("🔗 Risultato della ricerca web:")
        st.markdown(st.session_state["web_manual_input"])
    if col2.button("🧠 Deep Research", key="deep_manual_input"):
        st.session_state["deep_manual_input"] = perform_deep_research(manual_input)
        st.info("📌 Insight da deep research:")
        st.markdown(st.session_state["deep_manual_input"])

# Ranking & KPI
if st.button("📊 Mostra ranking & matrice KPI"):
    df = st.session_state.get("excel_df")
    if df is not None:
    ranked_df = analyze_triggers_and_rank(df, parsed_pdf, manual_input, buyer_personas)
    if not ranked_df.empty:
        st.subheader("🔎 Trigger → KPI → Messaggio suggerito")
        st.dataframe(ranked_df[[
            "Name", "Company", "Role", "Trigger combinato", 
            "Score", "KPI consigliati", "Framework suggerito", 
            "Note Deep"
        ]])
        st.markdown("\n")
        st.download_button(
            "📥 Scarica risultati", 
            data=ranked_df.to_csv(index=False).encode("utf-8"), 
            file_name="risultati_campagna.csv", 
            mime="text/csv"
        )
        st.session_state["ranked_df"] = ranked_df
        st.success("✔️ Analisi completata.")
    else:
        st.info("Nessun trigger noto trovato.")
else:
    st.warning("Carica prima un file Excel.")

# Step 3: Messaggi personalizzati
with st.expander("🧠 3. Genera messaggi personalizzati", expanded=False):
    if st.button("🧠 Genera messaggi personalizzati"):
        df = st.session_state.get("ranked_df")
        if df is not None:
            with st.spinner("🔄 Generazione messaggi in corso..."):
                selected_framework = st.selectbox("📐 Scegli un framework", ["Auto (da score)", "TIPPS", "TIPPS + COI", "Poke the Bear", "Harris NEAT"])
                output_df = generate_personalized_messages(df, framework_override=selected_framework)
                st.session_state["personalized_messages"] = output_df
                st.success("✔️ Messaggi generati con successo!")

output_df = st.session_state.get("personalized_messages")
if output_df is not None:
    st.subheader("📩 Messaggi personalizzati")
    for i, row in output_df.iterrows():
        st.markdown(f"##### 🎯 [{row['Nome']} – {row['Azienda']} – {row['Ruolo']}]")
        new_msg = st.text_area(f"✏️ Modifica messaggio", value=row['Messaggio generato'], key=f"msg_edit_{i}")
        output_df.at[i, "Messaggio generato"] = new_msg

    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        output_df.to_excel(writer, index=False, sheet_name="Messaggi")
    st.download_button("📥 Scarica messaggi", data=buffer.getvalue(), file_name="messaggi_personalizzati.xlsx")

    if st.button("💾 Salva tutti in libreria"):
        for _, row in output_df.iterrows():
            save_to_library("Messaggio personalizzato", row["Messaggio generato"])
        st.success("📚 Messaggi salvati nella libreria!")

# Step 4: Sequenza multicanale
with st.expander("🧩 4. Sequenza multicanale completa", expanded=False):
    df = st.session_state.get("personalized_messages")
    if df is not None and not df.empty:
        role_type = st.selectbox("👤 Tipo di ruolo", ["AE", "SDR"])
        include_inmail = st.checkbox("✉️ Includi InMail?", value=False)
        sequence_length = st.selectbox("🔢 Versione sequenza SDR", ["Standard (8 step)", "Lunga (11 step)"], index=0 if role_type == "AE" else 1)

        if st.button("🎯 Genera sequenza multicanale"):
            with st.spinner("🧠 Generazione sequenza in corso..."):
                sequence_df = generate_multichannel_sequence(
                    df,
                    role_type=role_type,
                    include_inmail=include_inmail,
                    sequence_type="long" if sequence_length == "Lunga (11 step)" else "standard"
                )
                st.session_state["multichannel_sequence"] = sequence_df
                st.success("✅ Sequenza multicanale generata!")

# Visualizzazione e opzioni su sequenza
sequence_df = st.session_state.get("multichannel_sequence")
if sequence_df is not None and not sequence_df.empty:
    st.subheader("📬 Sequenza multicanale generata")
    for i, row in sequence_df.iterrows():
        st.markdown(f"#### 🔹 Giorno {row['Giorno']} – {row['Tipo Azione']}")
        st.markdown(f"**Target:** {row['Nome']} ({row['Ruolo']} – {row['Azienda']})")
        new_msg = st.text_area("✏️ Modifica messaggio", value=row["Messaggio"], key=f"seq_msg_{i}")
        sequence_df.at[i, "Messaggio"] = new_msg

    show_chart = st.checkbox("📊 Mostra il calendario visivo della sequenza", value=True)
    if show_chart:
        st.markdown("### 🗓️ Calendario visuale (distribuzione per giorno)")
        action_filter = st.multiselect("Filtra per tipo azione", options=sequence_df["Tipo Azione"].unique(), default=sequence_df["Tipo Azione"].unique())
        target_filter = st.multiselect("Filtra per ruolo target", options=sequence_df["Ruolo"].unique(), default=sequence_df["Ruolo"].unique())

        filtered_df = sequence_df[
            (sequence_df["Tipo Azione"].isin(action_filter)) &
            (sequence_df["Ruolo"].isin(target_filter))
        ]

        day_counts = filtered_df["Giorno"].value_counts().sort_index()
        fig, ax = plt.subplots(figsize=(8, 3))
        ax.bar(day_counts.index, day_counts.values)
        ax.set_xlabel("Giorno")
        ax.set_ylabel("Numero azioni")
        ax.set_title("Distribuzione delle azioni nella sequenza")
        ax.grid(axis="y")
        st.pyplot(fig)

    st.markdown("### 🔁 Vuoi provare un altro framework?")
    new_framework = st.selectbox("📐 Scegli framework alternativo", ["TIPPS", "TIPPS + COI", "Poke the Bear", "Harris NEAT"])
    if st.button("♻️ Rigenera con framework alternativo"):
        with st.spinner("🧠 Rigenerazione in corso..."):
            alt_seq = generate_multichannel_sequence(
                df,
                role_type=role_type,
                include_inmail=include_inmail,
                sequence_type="long" if sequence_length == "Lunga (11 step)" else "standard",
                framework_override=new_framework
            )
            st.session_state["multichannel_sequence"] = alt_seq
            st.rerun()

    if st.button("📚 Salva tutta la sequenza in libreria"):
        for _, row in sequence_df.iterrows():
            save_to_library("Sequenza multicanale", row["Messaggio"])
        st.success("✅ Sequenza salvata nella libreria!")

    # Download CSV e Word
    csv_buffer = io.StringIO()
    sequence_df.to_csv(csv_buffer, index=False)
    st.download_button("📥 Scarica CSV della sequenza", data=csv_buffer.getvalue(), file_name="sequenza_multicanale.csv", mime="text/csv")

    doc = Document()
    doc.add_heading("Sequenza Multicanale", level=1)
    for _, row in sequence_df.iterrows():
        msg = f"Giorno {row['Giorno']} – {row['Tipo Azione']} – {row['Nome']} ({row['Ruolo']} – {row['Azienda']}):\n{row['Messaggio']}"
        doc.add_paragraph(msg)

    word_buffer = BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    st.download_button("📥 Scarica Word con messaggi", data=word_buffer, file_name="sequenza_multicanale.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Step 5: Affina con AI
with st.expander("🧠 5. Affina la sequenza con l’assistente AI", expanded=False):
    st.markdown("### 🤖 Chatta con l’assistente AI sui messaggi")
    user_prompt = st.text_area("Hai domande o vuoi modificarli con GPT?", key="chat_seq_prompt")

    if st.button("✉️ Invia alla chat AI", key="send_chat_seq_prompt"):
        if user_prompt:
            with st.spinner("💬 Risposta AI in corso..."):
                response = openai.ChatCompletion.create(
                    model="gpt-4o",
                    messages=[{"role": "user", "content": user_prompt}],
                    temperature=0.6,
                )
                st.markdown("**Risposta AI:**")
                st.write(response.choices[0].message.content)
        else:
            st.warning("Scrivi qualcosa prima di inviare.")

            st.write(response.choices[0].message.content)
    else:
        st.warning("Scrivi qualcosa prima di inviare.")


