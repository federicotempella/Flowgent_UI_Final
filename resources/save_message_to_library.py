# All'inizio di save_message_to_library.py:
import json
import os
import csv
from datetime import datetime
from docx import Document 

def save_message_to_library(
    nome,
    azienda,
    ruolo,
    framework,
    trigger,
    messaggio,
    note_deep,
    file_path="resources/messages_library.json"
):
    # Assicura che la directory esista
    os.makedirs(os.path.dirname(file_path), exist_ok=True)

    # Nuovo messaggio da salvare
    new_entry = {
        "timestamp": datetime.utcnow().isoformat() + "Z",
        "nome": nome,
        "azienda": azienda,
        "ruolo": ruolo,
        "framework": framework,
        "trigger": trigger,
        "messaggio": messaggio,
        "note_deep": note_deep
    }

    # Se il file esiste, carica il contenuto attuale
    if os.path.exists(file_path):
        with open(file_path, "r", encoding="utf-8") as file:
            try:
                data = json.load(file)
            except json.JSONDecodeError:
                data = []
    else:
        data = []

    # Aggiungi il nuovo messaggio
    data.append(new_entry)

    # Salva tutto nel file
    with open(file_path, "w", encoding="utf-8") as file:
        json.dump(data, file, indent=2, ensure_ascii=False)

    print(f"✅ Messaggio salvato in {file_path}")

# Funzione per esportare la libreria in formato CSV
def export_library_to_csv():
    """Legge messages_library.json ed esporta i messaggi in resources/messages_export.csv, aggiungendo la data odierna a ogni voce."""
    try:
        with open('resources/messages_library.json', 'r', encoding='utf-8') as f:
            data = json.load(f)
    except FileNotFoundError:
        print("Errore: file messages_library.json non trovato.")
        return

    # Ottenere la data corrente in formato GG/MM/AAAA
    data_oggi = datetime.now().strftime("%d/%m/%Y")

    # Aprire il file CSV per scrivere i dati
    with open('resources/messages_export.csv', 'w', encoding='utf-8', newline='') as csvfile:
        # Se la libreria è vuota, scriviamo solo l'intestazione (oppure nessun output)
        if not data:
            writer = csv.writer(csvfile)
            writer.writerow(["Messaggio", "Data"])
            # Libreria vuota: nessuna riga di messaggio da scrivere
            return

        # Se gli elementi sono dizionari, usiamo i loro campi come colonne
        if isinstance(data[0], dict):
            # Intestazioni CSV: campi esistenti + 'Data'
            fieldnames = list(data[0].keys()) + ["Data"]
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()
            for msg in data:
                # Copiamo ogni messaggio e aggiungiamo il campo Data
                riga = msg.copy()
                riga["Data"] = data_oggi
                writer.writerow(riga)
        else:
            # Se la libreria è una lista di testi semplici (non dizionari)
            writer = csv.writer(csvfile)
            writer.writerow(["Messaggio", "Data"])
            for msg in data:
                writer.writerow([msg, data_oggi])
    print(f"✅ Libreria esportata in CSV ({len(data)} messaggi).")

# Funzione per esportare la libreria in formato Word (.docx)
def export_library_to_word():
    """Legge messages_library.json ed esporta i messaggi in resources/messages_export.docx, aggiungendo la data odierna a ogni voce."""
    try:
        with open('resources/messages_library.json', 'r', encoding='utf-8') as f:
            data = json.load(f)
    except FileNotFoundError:
        print("Errore: file messages_library.json non trovato.")
        return

    data_oggi = datetime.now().strftime("%d/%m/%Y")
    doc = Document()

    # Se la libreria è vuota, inseriamo comunque un’intestazione base
    if not data:
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Messaggio"
        hdr_cells[1].text = "Data"
        # (nessun messaggio da aggiungere)
    else:
        if isinstance(data[0], dict):
            # Prepariamo la tabella con colonne per ogni campo + Data
            fieldnames = list(data[0].keys()) + ["Data"]
            cols = len(fieldnames)
            table = doc.add_table(rows=1, cols=cols)
            hdr_cells = table.rows[0].cells
            # Intestazioni della tabella
            for j, campo in enumerate(fieldnames):
                hdr_cells[j].text = str(campo)
            # Aggiungiamo una riga per ogni messaggio
            for msg in data:
                row_cells = table.add_row().cells
                # Scriviamo ogni valore del dizionario nelle colonne corrispondenti
                for j, campo in enumerate(fieldnames):
                    if campo == "Data":
                        row_cells[j].text = data_oggi
                    else:
                        row_cells[j].text = str(msg.get(campo, ""))
        else:
            # La libreria è una lista di stringhe (un solo campo "Messaggio")
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Messaggio"
            hdr_cells[1].text = "Data"
            for msg in data:
                row_cells = table.add_row().cells
                row_cells[0].text = str(msg)
                row_cells[1].text = data_oggi

    # Salviamo il documento Word
    doc.save('resources/messages_export.docx')
    print(f"✅ Libreria esportata in Word ({len(data)} messaggi).")


